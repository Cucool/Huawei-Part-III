from __future__ import annotations

import ast
import json
from collections import OrderedDict
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
LOW_RESOURCE_DIR = ROOT / "generated_low_resource_cases"
QA_REWRITE_PATH = ROOT / "generated_qa_rewrites" / "qa_local_model_rewrites.json"
OUTPUT_DIR = ROOT / "generated_expert_review_docs" / "simplified"


COUNTRIES = [
    {
        "key": "saudi",
        "country": "沙特",
        "folder": "沙特",
        "language_file": "arabic",
        "language_label": "阿拉伯语（沙特）",
    },
    {
        "key": "thailand",
        "country": "泰国",
        "folder": "泰国",
        "language_file": "thai",
        "language_label": "泰语（泰国）",
    },
    {
        "key": "turkey",
        "country": "土耳其",
        "folder": "土耳其",
        "language_file": "turkish",
        "language_label": "土耳其语（土耳其）",
    },
]


DIALECT_FILENAME = "部分方言化的规则及示例审核.docx"
REWRITE_FILENAME = "同义改写示例审核.docx"

DIALECT_RULE_SOURCES = {
    "arabic": (ROOT / "方言-沙特.py", "ARABIC_RULES"),
    "thai": (ROOT / "方言.py", "THAILAND_RULES"),
    "turkish": (ROOT / "方言.py", "TURKISH_RULES"),
}

TARGET_DIALECT_RULE_IDS = {
    "arabic": ["AR1", "AR2", "AR3", "AR4", "AR5", "AR6", "AR7", "AR8", "AR9", "AR10"],
    "thai": ["T1", "T2-1", "T2-2", "T3-1", "T3-2", "T3-3", "T3-4", "T4", "T5", "T6"],
    "turkish": ["TR1", "TR2-1", "TR2-2", "TR3", "TR4", "TR5-1", "TR5-2", "TR7", "TR9", "TR10"],
}

MISSING_RULE_EXAMPLES = {
    "arabic": {
        "AR6": {
            "original": "هذا بيت قديم.",
            "generated": "هذا بت قديم.",
        },
    },
}


def load_json(path: Path) -> list[dict[str, Any]]:
    with path.open("r", encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, list):
        raise ValueError(f"{path} 顶层必须为 list。")
    return [item for item in data if isinstance(item, dict)]


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).replace("\r\n", "\n").replace("\r", "\n").strip()


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(14)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    return doc


def add_text_paragraph(doc: Document, text: Any = ""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    lines = clean_text(text).splitlines()
    if not lines:
        return p
    p.add_run(lines[0])
    for line in lines[1:]:
        p.add_run().add_break(WD_BREAK.LINE)
        p.add_run(line)
    return p


def add_review_label(doc: Document, text: str):
    p = add_text_paragraph(doc, text)
    if p.runs:
        p.runs[0].bold = True
    return p


def add_title(doc: Document, title: str, country: dict[str, str]) -> None:
    p = add_text_paragraph(doc, f"《{title}》")
    for run in p.runs:
        run.bold = True
    add_text_paragraph(doc, f"审核对象：{country['country']} / {country['language_label']}")
    add_text_paragraph(doc)


def add_dialect_intro(doc: Document) -> None:
    add_text_paragraph(doc, "下面会给出一些把 原句子 部分方言化的规则及示例。")
    add_text_paragraph(doc, "需审核的内容为：")
    add_text_paragraph(doc, "（1）部分方言化的规则")
    add_text_paragraph(doc, "（2）部分方言化后的句子")
    add_text_paragraph(doc)
    add_text_paragraph(doc, "审核要求和注意事项为：")
    add_text_paragraph(doc, "1. 审核部分方言化的规则和部分方言化后的句子的合理性和正确性。")
    add_text_paragraph(
        doc,
        "2. 请注意这些规则和部分方言化后的句子的目的是为了将原句子部分方言化，不是将原句子完全转换为方言化的表达。",
    )
    add_text_paragraph(doc, "3. 审核部分方言化后的句子是否保留了原句子的意思。")
    p = add_text_paragraph(doc)
    p.add_run("4. 如果有需要修改，请以")
    p.add_run("修订模式").bold = True
    p.add_run("直接在 部分方言化的规则和部分方言化后的句子 上进行修改。")
    add_text_paragraph(doc)


def add_rewrite_intro(doc: Document) -> None:
    add_text_paragraph(doc, "下面会给出一些把 原句子 同义改写为 改写后的句子 的例子。")
    add_text_paragraph(doc, "需审核的内容为：")
    add_text_paragraph(doc, "（1）改写后的句子")
    add_text_paragraph(doc)
    add_text_paragraph(doc, "审核要求和注意事项为：")
    add_text_paragraph(doc, "1. 审核 改写后的句子 是否保留了 原句子 的意思。")
    add_text_paragraph(doc, "2. 审核 改写后的句子 的语法是否准确。")
    p = add_text_paragraph(doc)
    p.add_run("3. 如果有需要修改，请以")
    p.add_run("修订模式").bold = True
    p.add_run("直接在 改写后的句子 上进行修改。")
    add_text_paragraph(doc)


def collect_dialect_records(language: str) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    privacy = load_json(LOW_RESOURCE_DIR / f"{language}_privacy_dialect.json")
    safety = load_json(LOW_RESOURCE_DIR / f"{language}_safety_dialect_generalization.json")
    for item in privacy:
        record = dict(item)
        record["review_source_group"] = "隐私"
        records.append(record)
    for item in safety:
        record = dict(item)
        record["review_source_group"] = "内容安全"
        records.append(record)
    return records


def group_by_rule(records: list[dict[str, Any]]) -> OrderedDict[tuple[str, str], list[dict[str, Any]]]:
    grouped: OrderedDict[tuple[str, str], list[dict[str, Any]]] = OrderedDict()
    for item in records:
        key = (clean_text(item.get("rule_id")), clean_text(item.get("rule_description") or item.get("dialect_rule")))
        grouped.setdefault(key, []).append(item)
    return grouped


def load_rule_definitions(language: str) -> dict[str, str]:
    path, var_name = DIALECT_RULE_SOURCES[language]
    tree = ast.parse(path.read_text(encoding="utf-8"), filename=str(path))
    for node in tree.body:
        if not isinstance(node, ast.Assign):
            continue
        if not any(isinstance(target, ast.Name) and target.id == var_name for target in node.targets):
            continue
        rules = ast.literal_eval(node.value)
        return {
            clean_text(rule.get("id")): clean_text(rule.get("description"))
            for rule in rules
            if isinstance(rule, dict)
        }
    raise ValueError(f"无法在 {path} 中找到 {var_name}")


def selected_rule_groups(
    language: str,
    records: list[dict[str, Any]],
) -> list[tuple[str, str, list[dict[str, Any]]]]:
    by_rule_id: OrderedDict[str, list[dict[str, Any]]] = OrderedDict()
    for item in records:
        rule_id = clean_text(item.get("rule_id"))
        by_rule_id.setdefault(rule_id, []).append(item)

    definitions = load_rule_definitions(language)
    selected: list[tuple[str, str, list[dict[str, Any]]]] = []
    for rule_id in TARGET_DIALECT_RULE_IDS[language]:
        examples = by_rule_id.get(rule_id, [])
        rule_text = ""
        if examples:
            rule_text = clean_text(examples[0].get("rule_description") or examples[0].get("dialect_rule"))
        rule_text = rule_text or definitions.get(rule_id, "")
        if not examples and rule_id in MISSING_RULE_EXAMPLES.get(language, {}):
            fallback = dict(MISSING_RULE_EXAMPLES[language][rule_id])
            fallback.update(
                {
                    "rule_id": rule_id,
                    "rule_description": rule_text,
                    "dialect_rule": f"{rule_id}: {rule_text}",
                    "review_source_group": "规则补充示例",
                    "task_type_key": "",
                }
            )
            examples = [fallback]
        selected.append((rule_id, rule_text, examples))
    return selected


def split_mcq_parts(text: Any) -> list[tuple[str, str]]:
    raw = clean_text(text)
    if not raw:
        return []
    if "\n" in raw:
        pieces = [line.strip() for line in raw.splitlines() if line.strip()]
    else:
        pieces = [part.strip() for part in raw.split(" | ") if part.strip()]
    if not pieces:
        return []

    parts: list[tuple[str, str]] = [("题目", pieces[0])]
    for piece in pieces[1:]:
        label = piece.split(".", 1)[0].strip()
        if len(label) == 1 and label.isalpha() and "." in piece:
            parts.append((label.upper(), piece))
        else:
            parts.append(("", piece))
    return parts


def changed_mcq_text(original: Any, generated: Any) -> tuple[str, str]:
    original_parts = split_mcq_parts(original)
    generated_parts = split_mcq_parts(generated)
    if not original_parts or len(original_parts) != len(generated_parts):
        return clean_text(original), clean_text(generated)

    original_changed: list[str] = []
    generated_changed: list[str] = []
    for (original_label, original_text), (generated_label, generated_text) in zip(original_parts, generated_parts):
        if original_label != generated_label:
            return clean_text(original), clean_text(generated)
        if original_text == generated_text:
            continue
        if original_label == "题目":
            original_changed.append(f"题目：{original_text}")
            generated_changed.append(f"题目：{generated_text}")
        else:
            original_changed.append(original_text)
            generated_changed.append(generated_text)

    if not original_changed:
        return clean_text(original), clean_text(generated)
    return "\n".join(original_changed), "\n".join(generated_changed)


def dialect_display_texts(item: dict[str, Any]) -> tuple[str, str]:
    original = item.get("original")
    generated = item.get("generated")
    if item.get("task_type_key") == "knowledge":
        return changed_mcq_text(original, generated)
    return clean_text(original), clean_text(generated)


def add_dialect_records(doc: Document, records: list[dict[str, Any]], language: str) -> None:
    for rule_index, (rule_id, rule_text, examples) in enumerate(selected_rule_groups(language, records), start=1):
        add_review_label(doc, f"【需审核】规则{rule_index}：")
        add_text_paragraph(doc, rule_text or rule_id)
        add_text_paragraph(doc)
        add_text_paragraph(doc, f"规则{rule_index}的应用示例：")
        add_text_paragraph(doc)
        for example_index, item in enumerate(examples, start=1):
            original_text, generated_text = dialect_display_texts(item)
            add_text_paragraph(doc, f"原句子{example_index}：")
            add_text_paragraph(doc, original_text)
            add_review_label(doc, f"【需审核】部分方言化的句子{example_index}：")
            add_text_paragraph(doc, generated_text)
            add_text_paragraph(doc)


def collect_rewrite_records(country: dict[str, str]) -> list[dict[str, Any]]:
    language = country["language_file"]
    records: list[dict[str, Any]] = []
    for item in load_json(LOW_RESOURCE_DIR / f"{language}_privacy_qwen_rewrite.json"):
        record = dict(item)
        record["review_source_group"] = "隐私"
        records.append(record)

    for item in load_json(QA_REWRITE_PATH):
        if item.get("country") != country["country"]:
            continue
        record = dict(item)
        record["review_source_group"] = "内容安全"
        records.append(record)
    return records


def add_rewrite_records(doc: Document, records: list[dict[str, Any]]) -> None:
    for index, item in enumerate(records, start=1):
        add_text_paragraph(doc, f"原句子{index}：")
        add_text_paragraph(doc, item.get("original"))
        add_review_label(doc, f"【需审核】改写后的句子{index}：")
        add_text_paragraph(doc, item.get("rewrite") or item.get("case"))
        add_text_paragraph(doc)


def build_dialect_doc(country: dict[str, str]) -> Path:
    doc = setup_doc()
    add_title(doc, "部分方言化的规则及示例审核", country)
    add_dialect_intro(doc)
    add_dialect_records(doc, collect_dialect_records(country["language_file"]), country["language_file"])
    output_dir = OUTPUT_DIR / country["folder"]
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / DIALECT_FILENAME
    doc.save(output_path)
    return output_path


def build_rewrite_doc(country: dict[str, str]) -> Path:
    doc = setup_doc()
    add_title(doc, "同义改写示例审核", country)
    add_rewrite_intro(doc)
    add_rewrite_records(doc, collect_rewrite_records(country))
    output_dir = OUTPUT_DIR / country["folder"]
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / REWRITE_FILENAME
    doc.save(output_path)
    return output_path


def main() -> None:
    outputs: list[Path] = []
    for country in COUNTRIES:
        outputs.append(build_dialect_doc(country))
        outputs.append(build_rewrite_doc(country))
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
