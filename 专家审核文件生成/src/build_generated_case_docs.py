from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "generated_case_docs"

COUNTRIES = [
    {
        "key": "saudi",
        "country": "沙特",
        "language_file": "arabic",
        "language_label": "阿拉伯语（沙特）",
        "privacy_prefix": "沙特",
        "qa_prefix": "saudi",
        "output": "saudi_generated_cases.docx",
    },
    {
        "key": "thailand",
        "country": "泰国",
        "language_file": "thai",
        "language_label": "泰语（泰国）",
        "privacy_prefix": "泰国",
        "qa_prefix": "thailand",
        "output": "thailand_generated_cases.docx",
    },
    {
        "key": "turkey",
        "country": "土耳其",
        "language_file": "turkish",
        "language_label": "土耳其语（土耳其）",
        "privacy_prefix": "土耳其",
        "qa_prefix": "turkey",
        "output": "turkey_generated_cases.docx",
    },
]


def load_json(path: Path) -> list[dict[str, Any]]:
    with path.open("r", encoding="utf-8") as f:
        data = json.load(f)
    if isinstance(data, list):
        return [x for x in data if isinstance(x, dict)]
    raise ValueError(f"{path} must contain a JSON list")


LOW_RESOURCE_DIR = ROOT / "generated_low_resource_cases"
QA_REWRITE_PATH = ROOT / "generated_qa_rewrites" / "qa_local_model_rewrites.json"
JAILBREAK_PATH = ROOT / "jailbreak" / "generated" / "jailbreak_cases.json"

QA_REWRITES = load_json(QA_REWRITE_PATH)
JAILBREAK_CASES = load_json(JAILBREAK_PATH)
XML_INVALID_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    return XML_INVALID_RE.sub(" ", str(value))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_cell_margins(table, top=80, bottom=80, start=120, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    if "Case Meta" not in styles:
        style = styles.add_style("Case Meta", 1)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor.from_string("555555")
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.15

    if "Case Text" not in styles:
        style = styles.add_style("Case Text", 1)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(9.5)
        style.paragraph_format.left_indent = Inches(0.18)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    if "Record Title" not in styles:
        style = styles.add_style("Record Title", 1)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("1F4D78")
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.keep_with_next = True


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(subtitle)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("555555")


def add_footer(doc: Document, label: str) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.text = ""
    run = p.add_run(label + " | Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("555555")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    p._p.append(fld_begin)
    p._p.append(instr)
    p._p.append(fld_end)


def add_summary_table(doc: Document, rows: list[tuple[str, int, str]]) -> None:
    doc.add_heading("整理范围", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table)
    set_cell_margins(table)
    widths = [2800, 1200, 5360]
    headers = ["来源/类别", "数量", "说明"]
    for idx, cell in enumerate(table.rows[0].cells):
        cell.text = headers[idx]
        set_cell_shading(cell, "E8EEF5")
        set_cell_width(cell, widths[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    set_repeat_table_header(table.rows[0])

    for label, count, note in rows:
        cells = table.add_row().cells
        values = [label, str(count), note]
        for idx, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(9.5)


def add_label_text(doc: Document, label: str, text: Any, style: str = "Case Text") -> None:
    value = clean_text(text)
    p = doc.add_paragraph(style=style)
    r = p.add_run(label + "：")
    r.bold = True
    if "\n" not in value:
        p.add_run(value)
        return
    lines = value.splitlines()
    if lines:
        p.add_run(lines[0])
    for line in lines[1:]:
        p.add_run().add_break(WD_BREAK.LINE)
        p.add_run(line)


def add_record(doc: Document, title: str, meta: list[tuple[str, Any]], text_fields: list[tuple[str, Any]]) -> None:
    doc.add_paragraph(clean_text(title), style="Record Title")
    meta_text = " | ".join(f"{k}: {clean_text(v)}" for k, v in meta if v not in (None, ""))
    if meta_text:
        doc.add_paragraph(meta_text, style="Case Meta")
    for label, value in text_fields:
        add_label_text(doc, label, value)


def add_records_section(
    doc: Document,
    heading: str,
    records: list[dict[str, Any]],
    kind: str,
) -> None:
    doc.add_heading(f"{heading}（{len(records)} 条）", level=2)
    if not records:
        doc.add_paragraph("当前生成文件中没有对应记录。")
        return

    for idx, item in enumerate(records, start=1):
        if kind == "privacy_dialect":
            add_record(
                doc,
                f"{idx:03d} · 隐私模板 {item.get('template_index')} · {item.get('rule_id')}",
                [("方言规则", item.get("rule_description")), ("状态", item.get("status"))],
                [("original", item.get("original")), ("generated", item.get("generated"))],
            )
        elif kind == "safety_dialect":
            add_record(
                doc,
                f"{idx:03d} · {item.get('task_type')} · {item.get('rule_id')}",
                [
                    ("方言规则", item.get("rule_description")),
                    ("source_id", item.get("source_id")),
                    ("状态", item.get("status")),
                ],
                [("original", item.get("original")), ("generated", item.get("generated"))],
            )
        elif kind == "privacy_qwen":
            add_record(
                doc,
                f"{idx:03d} · 隐私模板 {item.get('template_index')}",
                [("model", item.get("model")), ("状态", item.get("status"))],
                [("original", item.get("original")), ("rewrite", item.get("rewrite"))],
            )
        elif kind == "qa_qwen":
            add_record(
                doc,
                f"{idx:03d} · {item.get('type')}",
                [("source_id", item.get("source_id")), ("model", item.get("model")), ("状态", item.get("status"))],
                [("original", item.get("original")), ("case", item.get("case"))],
            )
        elif kind == "jailbreak_privacy":
            add_record(
                doc,
                f"{idx:03d} · {item.get('method_id')} {item.get('method_name')}",
                [("source_item", item.get("source_item")), ("source_id", item.get("source_id"))],
                [("case", item.get("case")), ("jailbreak_case", item.get("jailbreak_case"))],
            )
        elif kind == "jailbreak_qa":
            rule_type = infer_rule_type(item.get("source_rule_id", ""))
            add_record(
                doc,
                f"{idx:03d} · {rule_type} · {item.get('method_id')} {item.get('method_name')}",
                [("rule_id", item.get("source_rule_id")), ("rule_zh", item.get("source_rule_zh"))],
                [("case", item.get("case")), ("jailbreak_case", item.get("jailbreak_case"))],
            )


def infer_rule_type(rule_id: str) -> str:
    if "_devaluation_" in rule_id:
        return "价值贬损"
    if "_violation_assistance_" in rule_id:
        return "违规协助"
    return "未知题型"


def country_records(config: dict[str, str]) -> dict[str, list[dict[str, Any]]]:
    lang = config["language_file"]
    country = config["country"]
    privacy_prefix = config["privacy_prefix"]
    qa_prefix = config["qa_prefix"]

    privacy_dialect = load_json(LOW_RESOURCE_DIR / f"{lang}_privacy_dialect.json")
    safety_dialect = load_json(LOW_RESOURCE_DIR / f"{lang}_safety_dialect.json")
    privacy_qwen = load_json(LOW_RESOURCE_DIR / f"{lang}_privacy_qwen_rewrite.json")
    qa_qwen = [x for x in QA_REWRITES if x.get("country") == country]
    jailbreak_privacy = [
        x
        for x in JAILBREAK_CASES
        if x.get("category") == "privacy" and str(x.get("source_item", "")).startswith(privacy_prefix + "_")
    ]
    jailbreak_qa = [
        x
        for x in JAILBREAK_CASES
        if x.get("category") == "QA" and str(x.get("source_rule_id", "")).startswith(qa_prefix + "_")
    ]

    return {
        "privacy_dialect": privacy_dialect,
        "safety_dialect": safety_dialect,
        "privacy_qwen": privacy_qwen,
        "qa_qwen": qa_qwen,
        "jailbreak_privacy": jailbreak_privacy,
        "jailbreak_qa": jailbreak_qa,
    }


def create_doc(config: dict[str, str]) -> Path:
    records = country_records(config)

    doc = Document()
    setup_styles(doc)
    add_footer(doc, f"{config['country']}生成用例整理")
    add_title(
        doc,
        f"{config['country']}生成用例整理",
        f"语言：{config['language_label']}；来源：generated_low_resource_cases、generated_qa_rewrites、jailbreak/generated",
    )

    rows = [
        ("方言化 - 隐私", len(records["privacy_dialect"]), "保留 original 与 generated；含模板序号、方言规则。"),
        ("方言化 - 内容安全", len(records["safety_dialect"]), "保留 original 与 generated；含题型、方言规则、source_id。"),
        ("本地大模型改写 - 隐私", len(records["privacy_qwen"]), "保留 original 与 rewrite。"),
        ("本地大模型改写 - 内容安全", len(records["qa_qwen"]), "保留 original 与 case；每个国家 3 个题型。"),
        ("15种越狱 - 隐私普通人名", len(records["jailbreak_privacy"]), "保留 case 与 jailbreak_case；含 source_item 和攻击方法。"),
        ("15种越狱 - 内容安全", len(records["jailbreak_qa"]), "保留 case 与 jailbreak_case；含 rule_id、rule_zh 和攻击方法。"),
    ]
    add_summary_table(doc, rows)

    doc.add_heading("1. 方言化", level=1)
    add_records_section(doc, "1.1 隐私", records["privacy_dialect"], "privacy_dialect")
    add_records_section(doc, "1.2 内容安全", records["safety_dialect"], "safety_dialect")

    doc.add_page_break()
    doc.add_heading("2. 本地大模型改写", level=1)
    add_records_section(doc, "2.1 隐私", records["privacy_qwen"], "privacy_qwen")
    add_records_section(doc, "2.2 内容安全", records["qa_qwen"], "qa_qwen")

    doc.add_page_break()
    doc.add_heading("3. 15种越狱", level=1)
    add_records_section(doc, "3.1 隐私：普通人名", records["jailbreak_privacy"], "jailbreak_privacy")
    add_records_section(doc, "3.2 内容安全", records["jailbreak_qa"], "jailbreak_qa")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUTPUT_DIR / config["output"]
    doc.save(out_path)
    return out_path


def main() -> None:
    output_paths = []
    for config in COUNTRIES:
        output_paths.append(create_doc(config))

    summary = {
        path.name: {
            "path": str(path),
            "size_bytes": path.stat().st_size,
        }
        for path in output_paths
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
